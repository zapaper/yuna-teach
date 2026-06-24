# Build the "Verb -> Noun" PSLE Synthesis one-pager in the brand
# format that matches MarkForYou-PSLE-English-Reported-Speech-v16.docx
#   - A4, 0.40" margins
#   - Title 24pt bold #1F2A37 Calibri
#   - Section header 12pt bold teal #0E6B6B
#   - Bullets 11pt bold teal
#   - Grid of tip-boxes: 14pt bold red header + 11pt grey body +
#     bold dark-green "after" line
#   - CTA 13pt bold teal at bottom
import io, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"c:\Users\peter\OneDrive\Documents\MarkForYou\MarkForYou-PSLE-English-Verb-To-Noun-v1.docx"

# Brand palette
SLATE   = RGBColor(0x1F, 0x2A, 0x37)
TEAL    = RGBColor(0x0E, 0x6B, 0x6B)
RED     = RGBColor(0xDC, 0x26, 0x26)
GREY    = RGBColor(0x6B, 0x72, 0x80)
GREEN   = RGBColor(0x06, 0x5F, 0x46)

doc = Document()

# Page setup — A4 with 0.40" margins (brand spec).
sec = doc.sections[0]
sec.page_width  = Inches(8.27)
sec.page_height = Inches(11.69)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(0.40)

def add(text, size, *, bold=False, color=None, after=0, align=None):
    p = doc.add_paragraph()
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p

def cell_para(cell, text, size, *, bold=False, color=None, after=2):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)

# ---- Header block ----
add("Verb-to-Noun: 1 in 8 PSLE Synthesis questions",
    26, bold=True, color=SLATE, after=2)
add("12 years of PSLE English Synthesis. When the keyword is made / "
    "came to / took / had a / the, you need this trick.",
    13, color=SLATE, after=4)

add("THE RULE", 14, bold=True, color=TEAL, after=1)
add("Turn the VERB (or adjective) in the original into a NOUN. "
    "Any -ly word that describes the verb (slightly, carefully) becomes an "
    "adjective (slight, careful).",
    13, bold=True, color=TEAL, after=1)
add("    adjusted slightly  →  made slight adjustments",
    13, bold=True, color=GREEN, after=4)

# ---- Two main "types" mini-table (Type A + Type B, side by side) ----
type_table = doc.add_table(rows=1, cols=2)
type_table.autofit = False
type_table.columns[0].width = Inches(3.73)
type_table.columns[1].width = Inches(3.73)

def emit_answer(cell, ans, highlight):
    """Render '✓ <ans>' with the substring `highlight` bolded and the
    rest in normal weight, all in green."""
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    idx = ans.find(highlight) if highlight else -1
    def green_run(text, *, bold):
        r = p.add_run(text)
        r.font.name = "Calibri"; r.font.size = Pt(11)
        r.bold = bold; r.font.color.rgb = GREEN
        return r
    green_run("✓ ", bold=True)
    if idx < 0:
        green_run(ans, bold=True)
        return
    green_run(ans[:idx], bold=False)
    green_run(highlight, bold=True)
    green_run(ans[idx + len(highlight):], bold=False)

def type_box(cell, label, pairs, examples):
    """examples: list of (year, source, answer, highlight) tuples."""
    set_cell_bg(cell, "F0FDFA")
    cell.width = Inches(3.73)
    cell_para(cell, label, 14, bold=True, color=TEAL, after=2)
    pair_rows = [(pairs[i], pairs[i+1] if i+1 < len(pairs) else None)
                 for i in range(0, len(pairs), 2)]
    for left, right in pair_rows:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        def emit(v, n):
            r1 = p.add_run(f"{v} → ")
            r1.font.name = "Calibri"; r1.font.size = Pt(12); r1.font.color.rgb = GREY
            r2 = p.add_run(n)
            r2.font.name = "Calibri"; r2.font.size = Pt(12); r2.bold = True; r2.font.color.rgb = SLATE
        emit(*left)
        if right:
            sep = p.add_run("     ")
            sep.font.name = "Calibri"; sep.font.size = Pt(12)
            emit(*right)
    cell_para(cell, "", 4, after=2)
    for year, src, ans, highlight in examples:
        cell_para(cell, year, 10, bold=True, color=GREY, after=0)
        cell_para(cell, src, 11, color=GREY, after=0)
        emit_answer(cell, ans, highlight)

type_box(
    type_table.rows[0].cells[0],
    "Type A — Noun form",
    [("lose", "loss"), ("conclude", "conclusion"),
     ("accurate", "accuracy"), ("humble", "humility")],
    [
        ("PSLE 2022    keyword: the    (kid converts: lost → the loss of)",
         "Vasanthi lost her mobile phone.",
         "Vasanthi reported the loss of her mobile phone to the police.",
         "the"),
        ("PSLE 2019    keyword: the winner's    (kid converts: humble → humility)",
         "The winner was humble. It moved the viewers.",
         "The viewers were moved by the winner's humility.",
         "the winner's"),
        ("PSLE 2021    keyword: The team discussed    (REVERSE: noun → verb)",
         "The discussion of the project took the team an hour.",
         "The team discussed the project for an entire hour.",
         "The team discussed"),
    ],
)
type_box(
    type_table.rows[0].cells[1],
    "Type B — Starter verb + noun",
    [("adjust", "made adjustments"), ("conclude", "came to a conclusion"),
     ("decide", "made a decision"), ("halt", "came to a halt")],
    [
        ("PSLE 2022    keyword: made",
         "The diver adjusted his goggles slightly.",
         "The diver made slight adjustments to his goggles.",
         "made"),
        ("PSLE 2022    keyword: came",
         "The meeting concluded only in the evening.",
         "The meeting came to a conclusion only in the evening.",
         "came"),
    ],
)

add("", 4, after=2)
add("COMMON MISTAKES", 14, bold=True, color=RED, after=2)

# ---- 4 tip boxes (2 x 2) ----
tip_table = doc.add_table(rows=2, cols=2)
tip_table.autofit = False
for col in tip_table.columns:
    col.width = Inches(3.73)

def tip_box(cell, n, headline, wrong, right):
    set_cell_bg(cell, "FEF2F2")
    cell.width = Inches(3.73)
    cell_para(cell, f"{n}.  {headline}", 14, bold=True, color=RED, after=3)
    cell_para(cell, f"✗  {wrong}", 13, bold=True, color=GREY, after=1)
    cell_para(cell, f"✓  {right}", 13, bold=True, color=GREEN, after=2)

tip_box(
    tip_table.rows[0].cells[0], 1,
    "Keep the describing word.",
    "adjusted slightly → made adjustments",
    "adjusted slightly → made slight adjustments",
)
tip_box(
    tip_table.rows[0].cells[1], 2,
    "Use the right starter-verb pair.",
    "did a decision  /  made a conclusion",
    "made a decision  /  came to a conclusion",
)
tip_box(
    tip_table.rows[1].cells[0], 3,
    "Use 's for the possessive noun.",
    "the winner humility",
    "the winner's humility",
)
tip_box(
    tip_table.rows[1].cells[1], 4,
    "Don't change the possessives.",
    "his goggles → the goggles",
    "his goggles → his goggles",
)

add("", 2, after=2)
add("Practice more Synthesis & Transformation at www.MarkForYou.com",
    13, bold=True, color=TEAL, after=0, align=WD_PARAGRAPH_ALIGNMENT.CENTER)

doc.save(OUT)
print(f"Saved: {OUT}")
